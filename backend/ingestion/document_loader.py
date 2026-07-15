"""Document loader with support for PDF, TXT, MD, and DOCX formats."""

from __future__ import annotations

import hashlib
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".txt", ".md", ".docx"})


@dataclass
class RawDocument:
    """A loaded document with extracted text and per-page content."""

    file_path: str
    filename: str
    text: str
    page_texts: list[tuple[int, str]] = field(default_factory=list)
    file_hash: str = ""


class DocumentLoader:
    """Scan directories and load documents from supported file types.

    Supports incremental indexing via SHA-256 file hashes.
    """

    def __init__(self, supported_extensions: frozenset[str] | None = None) -> None:
        self._extensions = supported_extensions or SUPPORTED_EXTENSIONS

    def scan_directory(self, directory: str | Path) -> list[Path]:
        """Recursively scan a directory for supported files.

        Args:
            directory: Path to the directory to scan.

        Returns:
            Sorted list of file paths with supported extensions.
        """
        root = Path(directory)
        if not root.is_dir():
            logger.warning("Directory does not exist: %s", root)
            return []

        files: list[Path] = []
        for path in sorted(root.rglob("*")):
            if path.is_file() and path.suffix.lower() in self._extensions:
                files.append(path)
        return files

    def load_document(self, file_path: Path) -> RawDocument:
        """Load a single document from disk.

        Args:
            file_path: Path to the document file.

        Returns:
            RawDocument with extracted text and file hash.

        Raises:
            RuntimeError: If the file cannot be read or parsed.
        """
        suffix = file_path.suffix.lower()
        file_hash = self._compute_hash(file_path)

        try:
            if suffix == ".pdf":
                text, page_texts = self._load_pdf(file_path)
            elif suffix in {".txt", ".md"}:
                text, page_texts = self._load_text(file_path)
            elif suffix == ".docx":
                text, page_texts = self._load_docx(file_path)
            else:
                raise RuntimeError(f"Unsupported file type: {suffix}")
        except Exception as exc:
            logger.error("Failed to load document %s: %s", file_path, exc)
            raise RuntimeError(f"Failed to load {file_path.name}: {exc}") from exc

        return RawDocument(
            file_path=str(file_path),
            filename=file_path.name,
            text=text,
            page_texts=page_texts,
            file_hash=file_hash,
        )

    def load_all(self, directory: str | Path) -> Iterator[RawDocument]:
        """Load all supported documents from a directory.

        Args:
            directory: Path to scan.

        Yields:
            RawDocument for each successfully loaded file.
        """
        files = self.scan_directory(directory)
        logger.info("Found %d supported files in %s", len(files), directory)

        for file_path in files:
            try:
                yield self.load_document(file_path)
            except RuntimeError as exc:
                logger.warning("Skipping %s: %s", file_path.name, exc)
                continue

    def _load_pdf(self, file_path: Path) -> tuple[str, list[tuple[int, str]]]:
        """Extract text from a PDF file page by page."""
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError(
                "pypdf is required for PDF ingestion. Install with: pip install pypdf"
            ) from exc

        reader = PdfReader(str(file_path))
        page_texts: list[tuple[int, str]] = []
        all_text_parts: list[str] = []

        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            page_text = page_text.strip()
            if page_text:
                page_texts.append((page_num, page_text))
                all_text_parts.append(page_text)

        return "\n\n".join(all_text_parts), page_texts

    def _load_text(self, file_path: Path) -> tuple[str, list[tuple[int, str]]]:
        """Read a plain text or markdown file as UTF-8."""
        text = file_path.read_text(encoding="utf-8", errors="ignore").strip()
        # For text files, treat entire content as page 1
        page_texts = [(1, text)] if text else []
        return text, page_texts

    def _load_docx(self, file_path: Path) -> tuple[str, list[tuple[int, str]]]:
        """Extract text from a DOCX file."""
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for DOCX ingestion. Install with: pip install python-docx"
            ) from exc

        doc = DocxDocument(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        # DOCX doesn't have page numbers; treat as single page
        page_texts = [(1, text)] if text else []
        return text, page_texts

    @staticmethod
    def _compute_hash(file_path: Path) -> str:
        """Compute SHA-256 hash of a file's contents for incremental indexing."""
        hasher = hashlib.sha256()
        with open(file_path, "rb") as f:
            for block in iter(lambda: f.read(8192), b""):
                hasher.update(block)
        return hasher.hexdigest()
